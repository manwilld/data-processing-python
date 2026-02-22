"""Test Report Generator — shake table seismic test report.

Builds a Word document matching Derek's test report format.
Opens the project template to inherit all styles, then builds
content programmatically.

Sections produced:
  1. Cover page
  2. Test Results Summary
  3. Test Procedure
  4. UUT Summary (one per UUT, with resonance plots)
  5. Seismic Run (one per run, with all plot sections)
  6. Appendix – Laboratory Accreditation
"""

import os
import re
import glob

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ───────────────────────────────────────────────────────────────────

def _bold_run(para, text, size_pt=None):
    run = para.add_run(text)
    run.bold = True
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def _set_cell_bold(cell, bold=True):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold


def _merge_header_row(table, row_idx, col_start, col_end, text, bold=True):
    """Merge cells in a row from col_start to col_end and set text."""
    row = table.rows[row_idx]
    cell = row.cells[col_start]
    for i in range(col_start + 1, col_end + 1):
        cell = cell.merge(row.cells[i])
    cell.text = text
    if bold:
        _set_cell_bold(cell)
    return cell


def _sorted_pngs(directory, pattern='*.png'):
    """Return sorted list of PNG paths in directory matching pattern."""
    if not directory or not os.path.isdir(directory):
        return []
    files = glob.glob(os.path.join(directory, pattern))
    return sorted(files)


def _filter_plots(directory, include=None, exclude=None):
    """Return sorted PNGs from directory, applying include/exclude substring filters."""
    paths = _sorted_pngs(directory)
    if include:
        paths = [p for p in paths if any(s in os.path.basename(p) for s in include)]
    if exclude:
        paths = [p for p in paths if not any(s in os.path.basename(p) for s in exclude)]
    return paths


def _select_photos(directory, max_photos=None,
                   extensions=('.jpg', '.jpeg', '.JPG', '.JPEG', '.png', '.PNG')):
    """Return a sorted list of photo paths from directory.

    If total count exceeds max_photos, returns an evenly-spaced sample that always
    includes the first and last photo. If max_photos is None, returns all.

    Parameters
    ----------
    directory  : str   Path to the photo directory.
    max_photos : int   Maximum number of photos to return (None = unlimited).
    extensions : tuple File extensions to include.
    """
    if not directory or not os.path.isdir(directory):
        return []
    photos = sorted([
        os.path.join(directory, f) for f in os.listdir(directory)
        if os.path.splitext(f)[1] in extensions
    ])
    if not photos:
        return []
    if max_photos is None or len(photos) <= max_photos:
        return photos
    if max_photos == 1:
        return [photos[len(photos) // 2]]
    # Evenly-spaced indices, always include first (0) and last (N-1)
    indices = {round(i * (len(photos) - 1) / (max_photos - 1)) for i in range(max_photos)}
    return [photos[i] for i in sorted(indices)]


# ── Generator ─────────────────────────────────────────────────────────────────

class TestReportGenerator:
    """Generate a complete shake table test report Word document."""

    PLOT_WIDTH = Inches(6.0)  # Full-width plot (fits within 1" margins)

    def __init__(self, resolved: dict):
        """
        Parameters
        ----------
        resolved : dict
            Merged project config (project_info + reports['test_report'] overrides).
        """
        self.cfg = resolved
        template_path = resolved.get('template')
        if template_path and os.path.isfile(template_path):
            self.doc = Document(template_path)
            self._clear_body()
        else:
            print(f"WARNING: Template not found at '{template_path}'. Using blank document.")
            self.doc = Document()
            self._setup_fallback_styles()

    # ── Document setup ────────────────────────────────────────────────────────

    def _clear_body(self):
        """Remove all existing paragraphs and tables from the template body."""
        body = self.doc.element.body
        for child in list(body):
            tag = child.tag.split('}')[-1]
            if tag in ('p', 'tbl', 'sdt'):
                body.remove(child)
        # Keep section properties (page size / margins)

    def _setup_fallback_styles(self):
        """Minimal style setup if template is missing."""
        for sec in self.doc.sections:
            sec.left_margin   = Inches(1.0)
            sec.right_margin  = Inches(1.0)
            sec.top_margin    = Inches(1.3)
            sec.bottom_margin = Inches(1.2)

    # ── Paragraph helpers ─────────────────────────────────────────────────────

    def _p(self, text='', style='Normal'):
        return self.doc.add_paragraph(text, style=style)

    def _h1(self, text):
        return self.doc.add_heading(text, level=1)

    def _h2(self, text):
        return self.doc.add_heading(text, level=2)

    def _no_space(self, text=''):
        return self._p(text, style='No Spacing')

    def _cover_header(self, text):
        return self._p(text, style='Cover Header')

    def _subheading(self, text):
        return self._p(text, style='SubHeading')

    def _tight(self, text):
        return self._p(text, style='Tight Spacing')

    def _blank(self):
        return self._no_space()

    # ── Plot embedding ────────────────────────────────────────────────────────

    def _embed_plot(self, png_path, centered=True):
        """Embed a PNG plot inline at full width."""
        if not png_path or not os.path.isfile(png_path):
            p = self._p(f'[Plot not found: {os.path.basename(png_path or "")}]')
            return p
        p = self.doc.add_paragraph()
        if centered:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(png_path, width=self.PLOT_WIDTH)
        return p

    def _embed_plots(self, png_paths, plots_per_continuation=4, continuation_heading=None):
        """Embed a list of plots, adding continuation headings every N plots."""
        for i, path in enumerate(png_paths):
            if continuation_heading and i > 0 and i % plots_per_continuation == 0:
                self._h2(continuation_heading)
            self._embed_plot(path)

    # ── Photo helpers ─────────────────────────────────────────────────────────

    @staticmethod
    def _resize_photo(src_path, max_px=900):
        """Return a BytesIO stream of the photo, resized to max_px on the longest side.

        Keeps aspect ratio. Returns None on error. Caches nothing — caller embeds directly.
        """
        try:
            from PIL import Image
            import io
            img = Image.open(src_path)
            # Convert to RGB if needed (e.g. CMYK or RGBA JPEGs)
            if img.mode not in ('RGB', 'L'):
                img = img.convert('RGB')
            w, h = img.size
            scale = min(max_px / max(w, h), 1.0)  # Never upscale
            if scale < 1.0:
                new_w = max(1, round(w * scale))
                new_h = max(1, round(h * scale))
                img = img.resize((new_w, new_h), Image.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format='JPEG', quality=80, optimize=True)
            buf.seek(0)
            return buf
        except Exception as e:
            print(f"  WARNING: could not resize {os.path.basename(src_path)}: {e}")
            return None

    def _build_photo_grid(self, photos, cols=2, photo_width_in=3.0, max_px=900):
        """Embed photos in a cols-wide grid table.

        Photos are resized to max_px on their longest side before embedding,
        keeping file size manageable.  Handles any number of photos: pads the
        last row with empty cells if the count is not a multiple of cols.

        Parameters
        ----------
        photos         : list[str]  Absolute paths to photo files.
        cols           : int        Number of columns (default 2).
        photo_width_in : float      Display width of each photo in inches.
        max_px         : int        Maximum pixel dimension before embedding (default 900).
        """
        if not photos:
            return None
        photo_w = Inches(photo_width_in)
        # Pad to a full row count
        padded = list(photos)
        remainder = len(padded) % cols
        if remainder:
            padded += [None] * (cols - remainder)
        n_rows = len(padded) // cols

        t = self._add_table(n_rows, cols, style='Table Grid')
        for ri in range(n_rows):
            for ci in range(cols):
                photo = padded[ri * cols + ci]
                cell = t.rows[ri].cells[ci]
                if photo and os.path.isfile(photo):
                    buf = self._resize_photo(photo, max_px=max_px)
                    if buf:
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.add_run().add_picture(buf, width=photo_w)
                # else: leave cell empty
        return t

    # ── Table helpers ─────────────────────────────────────────────────────────

    def _add_table(self, rows, cols, style='Table Grid'):
        t = self.doc.add_table(rows=rows, cols=cols)
        t.style = style
        return t

    def _set_row_bold(self, row):
        for cell in row.cells:
            _set_cell_bold(cell)

    def _fill_row(self, row, values, bold=False):
        for cell, val in zip(row.cells, values):
            cell.text = str(val)
            if bold:
                _set_cell_bold(cell)

    # ── TRS Excel reader ──────────────────────────────────────────────────────

    def _read_trs_excel(self, excel_path):
        """Read TRS table from Excel. Returns list of row dicts."""
        try:
            import openpyxl
        except ImportError:
            print("WARNING: openpyxl not installed — TRS table will be empty.")
            return [], {}

        if not excel_path or not os.path.isfile(excel_path):
            print(f"WARNING: TRS Excel not found: {excel_path}")
            return [], {}

        wb = openpyxl.load_workbook(excel_path, data_only=True)
        ws = wb.active

        # Read annotations from row 3 and 4 (low resonance, cutoff)
        annotations = {
            'low_resonance': ws.cell(3, 11).value,
            'low_resonance_label': ws.cell(3, 12).value,
            'cutoff': ws.cell(4, 11).value,
            'cutoff_label': ws.cell(4, 12).value,
        }

        # Data starts at row 3 (row 1=direction headers, row 2=column headers)
        # Columns: X(0-2), Y(3-5), Z(6-8)  → Freq, RRS, TRS for each direction
        rows = []
        for row in ws.iter_rows(min_row=3, values_only=True):
            if row[0] is None:
                continue
            rows.append({
                'freq_x': row[0], 'rrs_x': row[1], 'trs_x': row[2],
                'freq_y': row[3], 'rrs_y': row[4], 'trs_y': row[5],
                'freq_z': row[6], 'rrs_z': row[7], 'trs_z': row[8],
            })
        return rows, annotations

    def _build_trs_table(self, data_rows):
        """Build the Response Spectra Data table (3 direction groups × 3 cols)."""
        if not data_rows:
            self._p('[TRS data not available]')
            return

        n_data = len(data_rows)
        # 2 header rows + data rows
        t = self._add_table(2 + n_data, 9)

        # Row 0: direction headers (merged 3 cols each)
        _merge_header_row(t, 0, 0, 2, 'X Direction', bold=True)
        _merge_header_row(t, 0, 3, 5, 'Y Direction', bold=True)
        _merge_header_row(t, 0, 6, 8, 'Z Direction', bold=True)

        # Row 1: column headers
        col_headers = [
            'Freq.\n(Hz)', 'RRS\n(g)', 'TRS\n(g)',
            'Freq.\n(Hz)', 'RRS\n(g)', 'TRS\n(g)',
            'Freq.\n(Hz)', 'RRS\n(g)', 'TRS\n(g)',
        ]
        self._fill_row(t.rows[1], col_headers, bold=True)

        # Data rows
        for i, row in enumerate(data_rows):
            def _fmt(v):
                if v is None:
                    return ''
                try:
                    return f'{float(v):.2f}'
                except (TypeError, ValueError):
                    return str(v)

            vals = [
                _fmt(row['freq_x']), _fmt(row['rrs_x']), _fmt(row['trs_x']),
                _fmt(row['freq_y']), _fmt(row['rrs_y']), _fmt(row['trs_y']),
                _fmt(row['freq_z']), _fmt(row['rrs_z']), _fmt(row['trs_z']),
            ]
            self._fill_row(t.rows[2 + i], vals)

    def _build_levels_table(self, levels):
        """Seismic levels table (10 cols, 2 data rows per level: z/h=1 and z/h=0).
        Reused in Test Results Summary and at the start of each Run section.
        """
        n_data_rows = len(levels) * 2
        t = self._add_table(1 + n_data_rows, 10)
        self._fill_row(t.rows[0],
            ['Level', 'SDS (g)', 'z/h', 'Hf / Rμ',
             'AFLX-H (g)', 'ARIG-H (g)', 'AFLX-V (g)', 'ARIG-V (g)',
             '0.9·ARIG-H', '0.9·ARIG-V'],
            bold=True)
        row_idx = 1
        for lv in levels:
            hf_rmu_zh1 = f"{lv['Hf']} / {lv['Rmu']}"
            hf_rmu_zh0 = f"{lv.get('Hf_zh0', 1.0)} / {lv.get('Rmu_zh0', 1.0)}"
            self._fill_row(t.rows[row_idx], [
                lv['name'], f"{lv['Sds_zh1']:.2f}", '1', hf_rmu_zh1,
                f"{lv['Aflx_h']:.2f}", f"{lv['Arig_h']:.2f}",
                f"{lv['Aflx_v']:.2f}", f"{lv['Arig_v']:.2f}",
                f"{lv['Arig_h_90']:.2f}", f"{lv['Arig_v_90']:.2f}",
            ])
            self._fill_row(t.rows[row_idx + 1], [
                lv['name'], f"{lv['Sds_zh0']:.2f}", '0', hf_rmu_zh0,
                f"{lv['Aflx_h']:.2f}", f"{lv['Arig_h']:.2f}",
                f"{lv['Aflx_v']:.2f}", f"{lv['Arig_v']:.2f}",
                f"{lv['Arig_h_90']:.2f}", f"{lv['Arig_v_90']:.2f}",
            ])
            row_idx += 2
        return t

    def _build_run_results_table(self, runs, levels):
        """Peak acceleration results table for a list of runs.
        Reused in Test Results Summary and at the end of each Run section.
        """
        has_diag = any(r.get('has_diagonal', False) for r in runs)
        n_accel_cols = 4 if has_diag else 3
        n_cols = 5 + n_accel_cols
        accel_header = ['X', 'Y', '45', 'Z'] if has_diag else ['X', 'Y', 'Z']

        t = self._add_table(2 + len(runs), n_cols)
        row0_vals = (['Test Run', 'Test Date', 'Level', '0.9·ARIG-H', '0.9·ARIG-V']
                     + ['Peak Table Accel. (g)'] * n_accel_cols)
        self._fill_row(t.rows[0], row0_vals, bold=True)
        _merge_header_row(t, 0, 5, 4 + n_accel_cols, 'Peak Table Accel. (g)', bold=True)
        self._fill_row(t.rows[1],
            ['Test Run', 'Test Date', 'Level', '0.9·ARIG-H', '0.9·ARIG-V'] + accel_header,
            bold=True)
        for i, run in enumerate(runs):
            lv_name = run.get('level', '')
            lv = next((l for l in levels if l['name'] == lv_name), {})
            pa = run.get('peak_accel', {})
            accel_vals = [str(pa.get('X', '')), str(pa.get('Y', ''))]
            if has_diag:
                accel_vals.append(str(pa.get('D45', '')))
            accel_vals.append(str(pa.get('Z', '')))
            self._fill_row(t.rows[2 + i], [
                run.get('name', ''), run.get('date', ''), lv_name,
                f"{lv.get('Arig_h_90', '')}", f"{lv.get('Arig_v_90', '')}",
            ] + accel_vals)
        return t

    # ═════════════════════════════════════════════════════════════════════════
    # SECTION BUILDERS
    # ═════════════════════════════════════════════════════════════════════════

    def _build_cover(self):
        cfg = self.cfg

        # Title and subtitle
        self._p(cfg.get('title', ''), style='Title')
        self._p(cfg.get('subtitle', ''), style='Subtitle')

        # Testing scope
        self._cover_header('Testing Scope')
        self._p(cfg.get('testing_scope', ''))

        # Test units table (Manufacturer | Testing Laboratory)
        self._cover_header('Test Units')
        t = self._add_table(2, 2)
        t.rows[0].cells[0].text = 'Manufacturer'
        t.rows[0].cells[1].text = 'Testing Laboratory'
        self._set_row_bold(t.rows[0])
        mfr = cfg.get('manufacturer', {})
        lab = cfg.get('lab', {})
        t.rows[1].cells[0].text = (
            f"{mfr.get('company', '')}\n"
            f"{mfr.get('address', '')}\n"
            f"Contact: {mfr.get('contact', '')}"
        )
        t.rows[1].cells[1].text = (
            f"{lab.get('name', '')}\n"
            f"{lab.get('address', '')}\n"
            f"Contact: {lab.get('contact', '')}\n"
            f"Table: {lab.get('table_short', '')}"
        )

        # Certification company
        self._blank()
        self._cover_header('Certification company')
        eng = cfg.get('engineer', {})
        self._tight(eng.get('company', ''))
        self._tight(f"Certifying Engineer: {eng.get('name', '')}")
        self._tight(eng.get('license', ''))
        for _ in range(3):
            self._blank()

        # Revision history
        self._cover_header('REVISION HISTORY')
        revisions = cfg.get('revision_history', [])
        t = self._add_table(1 + len(revisions), 3)
        self._fill_row(t.rows[0], ['Revision', 'Date', 'Revision Description'], bold=True)
        for i, rev in enumerate(revisions):
            self._fill_row(t.rows[1 + i], [rev.get('rev',''), rev.get('date',''), rev.get('description','')])

    def _build_test_results_summary(self):
        cfg = self.cfg
        levels = cfg.get('levels', [])
        uuts   = cfg.get('uuts', [])
        runs   = cfg.get('runs', [])

        self._h1('Test Results Summary')

        # Intro — separate paragraphs matching final report structure
        self._p(f"Testing was successfully performed on the {cfg.get('subtitle', '')}.")
        for lv in levels:
            self._p(f"Level {lv['name']}: SDS={lv['Sds_zh1']:.2f}g (z/h=1); SDS={lv['Sds_zh0']:.2f}g (z/h=0)")
        self._blank()
        self._no_space(
            'As shown below, the acceleration time histories met the 90% ARIG requirement '
            'and the test was completed successfully.'
        )
        for _ in range(3):
            self._blank()

        # ── Seismic levels table ───────────────────────────────────────────────
        self._build_levels_table(levels)
        self._blank()

        # ── UUT summary table (7 cols, merged Dimensions header) ──────────────
        t = self._add_table(2 + len(uuts), 7)
        # Row 0: merged header for Dimensions
        self._fill_row(t.rows[0], ['UUT', 'Model / Description', 'Mounting', 'Dimensions (in)', '', '', 'Weight\n(lb)'], bold=True)
        _merge_header_row(t, 0, 3, 5, 'Dimensions (in)', bold=True)
        # Row 1: sub-headers
        self._fill_row(t.rows[1], ['UUT', 'Model / Description', 'Mounting', 'Depth', 'Width', 'Height', 'Weight\n(lb)'], bold=True)
        for i, u in enumerate(uuts):
            w = u.get('weight', '')
            w_str = f'{int(w):,}' if w != '' else ''
            self._fill_row(t.rows[2 + i], [
                str(u['number']), u['model'], u.get('mounting', ''),
                str(u.get('depth', '')), str(u.get('width', '')), str(u.get('height', '')),
                w_str,
            ])
        self._blank()

        # ── Resonance results table (7 cols, merged Resonant Freq header) ──────
        t = self._add_table(2 + len(uuts), 7)
        self._fill_row(t.rows[0], ['UUT', 'Test Date', 'Level', 'Result', 'Resonant Freq. (Hz)', '', ''], bold=True)
        _merge_header_row(t, 0, 4, 6, 'Resonant Freq. (Hz)', bold=True)
        self._fill_row(t.rows[1], ['UUT', 'Test Date', 'Level', 'Result', 'F-B', 'S-S', 'V'], bold=True)
        for i, u in enumerate(uuts):
            nf = u.get('nat_freq', {})
            self._fill_row(t.rows[2 + i], [
                str(u['number']), u.get('test_date', ''), u.get('level', ''),
                u.get('result', ''),
                str(nf.get('fb', '')), str(nf.get('ss', '')), str(nf.get('v', '')),
            ])
        self._blank()

        # ── Seismic run results table ─────────────────────────────────────────
        self._build_run_results_table(runs, levels)
        self._blank()

        # ── Lab equipment table ───────────────────────────────────────────────
        equip = cfg.get('lab_equipment', [])
        t = self._add_table(1 + len(equip), 8)
        self._fill_row(t.rows[0],
            ['Lab ID', 'Ch.', 'Description', 'Manufacturer', 'Model', 'Serial', 'Cal. Date', 'Cal. Due'],
            bold=True)
        for i, eq in enumerate(equip):
            self._fill_row(t.rows[1 + i], [
                eq.get('lab_id',''), eq.get('ch',''), eq.get('description',''),
                eq.get('manufacturer',''), eq.get('model',''), eq.get('serial',''),
                eq.get('cal_date',''), eq.get('cal_due',''),
            ])

    def _build_test_procedure(self):
        cfg = self.cfg
        eng = cfg.get('engineer', {})
        lab = cfg.get('lab', {})
        mfr = cfg.get('manufacturer', {})

        self._h1('Test Procedure')
        self._p(
            f"The following test procedure was implemented in accordance with "
            f"{cfg.get('test_standard', '')}. "
            f"Detailed photographs and plots from the seismic test are included in the "
            f"test run sections."
        )

        self._subheading('Pre-Test Inspection, Weighing, and Measuring')
        self._p(
            'Upon arrival at the lab, each UUT was visually examined for structural damage '
            'from transportation. The UUT weight and dimensions were measured and recorded.'
        )

        self._subheading('Pre-Test Functional Verification')
        self._p(
            'Functional testing was performed by the manufacturer at the test laboratory '
            'prior to performing seismic testing.'
        )

        self._subheading('Mounting')
        self._p(
            'Each UUT was mounted to the test fixture as described in the UUT Summary. '
            'The fixture was then attached to the shake table.'
        )

        self._subheading('Monitoring')
        self._p(
            'Reference control accelerometers were located on the shake table. '
            'These accelerometers were used as the control channels for the seismic test. '
            'Additional accelerometers were placed on each UUT to monitor responses.'
        )

        self._subheading('Resonant Frequency Search')
        self._p(
            'A resonant frequency search was performed prior to performing the '
            'multi-frequency seismic simulation test. The resonant frequency search '
            'consisted of a swept-sine excitation from 1 Hz to 33 Hz in each '
            'orthogonal direction.'
        )

        self._subheading('Multi-Frequency Seismic Simulation Test')
        self._p(
            f"Each UUT was subjected to a multi-frequency seismic simulation test "
            f"with the parameters shown in the Test Results Summary. The test was "
            f"performed in accordance with {cfg.get('test_standard', '')}."
        )

        self._subheading('Post-Test Inspection')
        self._p(
            'After the multi-frequency seismic simulation test, each UUT was visually '
            'examined. No visible structural damage was observed.'
        )

        self._subheading('Post-Test Functional Verification')
        self._p(
            'Functional testing was performed by the manufacturer at the test laboratory '
            'after performing seismic testing.'
        )

        # Witnesses
        self._cover_header('Test Witnesses')
        self._tight(mfr.get('witnesses', ''))
        self._tight(lab.get('witnesses', ''))
        self._tight(eng.get('witnesses', ''))

        # Shake table info
        self._cover_header('Shake Table Information')
        self._p(f"Table: {lab.get('table_long', '')}")
        self._p(lab.get('accreditation', ''))

    def _build_uut_summary(self, uut):
        cfg = self.cfg
        n   = uut['number']
        nf  = uut.get('nat_freq', {})
        lvs = cfg.get('levels', [])
        lv  = next((l for l in lvs if l['name'] == uut.get('level', '1')), lvs[0] if lvs else {})

        self._h1(f"UUT Summary \u2013 UUT {n}")

        # Basic info block
        self._no_space(f"Manufacturer: {uut.get('manufacturer', '')} \u2013 Model: {uut.get('model', '')}")
        self._p(f"UUT Function: {uut.get('function', '')};  Serial/Unique ID: {uut.get('serial', '')}")
        w = uut.get('weight', '')
        w_str = f'{int(w):,}' if w != '' else ''
        self._p(f"Depth = {uut.get('depth','')}-in, Width = {uut.get('width','')}-in, "
                f"Height = {uut.get('height','')}-in, Weight = {w_str}-lb")
        if lv:
            self._p(f"Maximum Test Level: SDS={lv.get('Sds_zh1',''):.2f}-g (z/h=1); "
                    f"SDS={lv.get('Sds_zh0',''):.2f}-g (z/h=0)")
        self._p(f"Resonant Frequencies: F-B = {nf.get('fb','')}-Hz, "
                f"S-S = {nf.get('ss','')}-Hz, V = {nf.get('v','')}-Hz")
        self._no_space('Mounting:')
        self._p(uut.get('mounting', ''))
        self._blank()

        # Continued section
        self._h2(f"UUT SUMMARY - UUT {n} (continued)")
        self._subheading('Construction:')
        self._p(uut.get('construction', ''))
        self._subheading('Subcomponents:')
        self._p(uut.get('subcomponents', ''))
        self._subheading('Function Testing:')
        self._p(uut.get('function_testing', ''))
        self._subheading('Testing Notes & Anomalies:')
        self._p(uut.get('notes', ''))

        # Resonance plots
        res_dirs = cfg.get('resonance_dirs', {})
        res_dir  = res_dirs.get(n)
        res_plots = _sorted_pngs(res_dir) if res_dir else []

        self._h2(f"RESONANT FREQUENCY SEARCH PLOTS - UUT {n}")
        per_page = 3
        for i, path in enumerate(res_plots):
            if i > 0 and i % per_page == 0:
                self._h2(f"RESONANT FREQUENCY SEARCH PLOTS - UUT {n} (continued)")
            self._embed_plot(path)

        if not res_plots:
            self._p(f'[No resonance plots found in: {res_dir}]')

    def _build_seismic_run(self, run):
        cfg    = self.cfg
        levels = cfg.get('levels', [])
        name   = run.get('name', 'Run')
        lv_name = run.get('level', '')
        lv = next((l for l in levels if l['name'] == lv_name), {})

        # Get paths from resolved config
        run_plots_cfg = cfg.get('run_plots', {}).get(name, {})
        seismic_dir   = run_plots_cfg.get('seismic_dir')
        trs_excel     = run_plots_cfg.get('trs_excel')

        self._h1(f"Seismic Run - {name}")
        self._no_space(
            f"The seismic level for {name} is shown below. "
            f"Detailed photographs and plots from the seismic test are included in "
            f"the sections below."
        )
        self._blank()

        # ── Seismic level table for this run ──────────────────────────────────
        # Show only the level(s) applicable to this run
        run_levels = [l for l in levels if l['name'] == lv_name] or levels
        self._build_levels_table(run_levels)
        self._blank()

        # ── Pre-Test Pictures ─────────────────────────────────────────────────
        self._h2(f"Pre-Test Pictures ({name})")
        pre_cfg = run_plots_cfg.get('pre_test_photos', {})
        pre_dir = pre_cfg.get('dir') if isinstance(pre_cfg, dict) else pre_cfg
        if pre_dir:
            pre_photos = _select_photos(pre_dir, max_photos=pre_cfg.get('max') if isinstance(pre_cfg, dict) else None)
            if pre_photos:
                n_pre = len(pre_photos)
                total_pre = len(_select_photos(pre_dir))
                print(f"  Embedding {n_pre} pre-test photos (of {total_pre} total)")
                self._build_photo_grid(pre_photos, cols=2, photo_width_in=3.0)
            else:
                self._p(f'[No pre-test photos found in: {pre_dir}]')
        else:
            self._p('[Pre-test photographs — see project photo archive]')

        # ── Post-Test Pictures ────────────────────────────────────────────────
        self._h2(f"Post-Test Pictures ({name})")
        post_cfg = run_plots_cfg.get('post_test_photos', {})
        post_dir = post_cfg.get('dir') if isinstance(post_cfg, dict) else post_cfg
        if post_dir:
            post_photos = _select_photos(post_dir, max_photos=post_cfg.get('max') if isinstance(post_cfg, dict) else None)
            if post_photos:
                n_post = len(post_photos)
                total_post = len(_select_photos(post_dir))
                print(f"  Embedding {n_post} post-test photos (of {total_post} total)")
                self._build_photo_grid(post_photos, cols=2, photo_width_in=3.0)
            else:
                self._p(f'[No post-test photos found in: {post_dir}]')
        else:
            self._p('[Post-test photographs — see project photo archive]')

        # ── Response Spectra Plots ────────────────────────────────────────────
        self._h2(f"Response Spectra Plots ({name})")
        # TRS plots: Table channel TRS + TRSall
        trs_plots = _filter_plots(
            seismic_dir,
            include=['TRS', 'TRSvsRRS'],
            exclude=['TH_', 'CC', 'CH'],
        )
        if trs_plots:
            for p in trs_plots:
                self._embed_plot(p)
        else:
            self._p(f'[No TRS plots found in: {seismic_dir}]')

        # ── Response Spectra Data (TRS table from Excel) ──────────────────────
        self._h2(f"Response Spectra Data ({name})")
        # Lowest resonance note
        if lv:
            uuts = cfg.get('uuts', [])
            all_nat = [min(u.get('nat_freq',{}).values()) for u in uuts if u.get('nat_freq')]
            if all_nat:
                low_res = min(all_nat)
                self._p(
                    f"As shown in the resonance search plots, the lowest resonant "
                    f"frequency is {low_res:.1f} Hz. The low cutoff frequency is "
                    f"determined from the resonance search."
                )
        trs_rows, trs_annotations = self._read_trs_excel(trs_excel)
        self._build_trs_table(trs_rows)

        # ── Acceleration Time History Plots ───────────────────────────────────
        self._h2(f"Acceleration Time History Plots ({name})")
        self._p(
            f"Per {cfg.get('test_standard', '')} Section 6.5.4.2.3, the peak shake table "
            f"acceleration shall equal or exceed 90 percent of ARIG in each orthogonal direction."
        )
        th_plots = _filter_plots(seismic_dir, include=['TH_Table'])
        for p in th_plots:
            self._embed_plot(p)
        if not th_plots:
            self._p(f'[No Table TH plots found in: {seismic_dir}]')

        # ── Statistical Independence Plots ────────────────────────────────────
        self._h2(f"Statistical Independence Plots ({name})")
        self._p(
            f"Per {cfg.get('test_standard', '')} Section 5.2.2.7.3, simultaneous shake table "
            f"motion in three orthogonal directions must be statistically independent."
        )
        cc_plots = _filter_plots(seismic_dir, include=['_CC', '_CH'])
        for p in cc_plots:
            self._embed_plot(p)
        if not cc_plots:
            self._p(f'[No CC/CH plots found in: {seismic_dir}]')

        # ── Unit Accelerometer Plots ──────────────────────────────────────────
        self._h2(f"Unit Accelerometer Plots ({name})")
        # UUT plots: everything that's not a Table channel, TRSall, CC, or CH
        uut_plots = _filter_plots(
            seismic_dir,
            exclude=['Table_', 'TRSvsRRS_All', '_CC', '_CH'],
        )
        plots_per_continuation = 4
        for i, path in enumerate(uut_plots):
            if i > 0 and i % plots_per_continuation == 0:
                self._h2(f"Unit Accelerometer Plots ({name}) (continued)")
            self._embed_plot(path)
        if not uut_plots:
            self._p(f'[No UUT accelerometer plots found in: {seismic_dir}]')

        # ── Run results summary table (repeated at end of each run) ───────────
        self._blank()
        self._build_run_results_table([run], levels)

    def _build_appendix(self):
        cfg = self.cfg
        lab = cfg.get('lab', {})
        self._h1('Appendix \u2013 Laboratory Accreditation')
        self._p(
            f"{lab.get('accreditation', '')} "
            f"A copy of the laboratory accreditation certificate is attached."
        )

    # ═════════════════════════════════════════════════════════════════════════
    # MAIN GENERATE
    # ═════════════════════════════════════════════════════════════════════════

    def generate(self, output_path: str) -> str:
        """Build and save the complete test report.

        Parameters
        ----------
        output_path : str
            Where to save the .docx file.

        Returns
        -------
        str
            Absolute path of the saved file.
        """
        print("Building cover page...")
        self._build_cover()

        print("Building Test Results Summary...")
        self.doc.add_page_break()
        self._build_test_results_summary()

        print("Building Test Procedure...")
        self.doc.add_page_break()
        self._build_test_procedure()

        for uut in self.cfg.get('uuts', []):
            print(f"Building UUT Summary – UUT {uut['number']}...")
            self.doc.add_page_break()
            self._build_uut_summary(uut)

        for run in self.cfg.get('runs', []):
            print(f"Building Seismic Run – {run.get('name', '')}...")
            self.doc.add_page_break()
            self._build_seismic_run(run)

        print("Building Appendix...")
        self.doc.add_page_break()
        self._build_appendix()

        output_path = os.path.abspath(output_path)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        print(f"Saved: {output_path}")
        return output_path
